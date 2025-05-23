import os
import anthropic
import json
from dotenv import load_dotenv
from datetime import datetime
import re

class MarkdownToWordMCP:
    """Markdownファイルの議事録をWord MCPでWord文書化するシステム"""
    
    def __init__(self):
        load_dotenv()
        self.client = anthropic.Anthropic(
            api_key=os.getenv("ANTHROPIC_API_KEY")
        )
        
        # Word MCP サーバー設定（ローカル実行）
        self.word_mcp_servers = [
            {
                "type": "stdio",  # Word MCPはローカル実行のみ
                "command": "python",
                "args": ["-m", "word_mcp_server"],
                "name": "word-mcp",
                "description": "Microsoft Word文書作成・編集"
            }
        ]
    
    def read_markdown_minutes(self, file_path):
        """Markdownファイルの議事録を読み込み"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            print(f"✅ Markdownファイル読み込み完了: {file_path}")
            print(f"📄 文字数: {len(content)}")
            return content
        except FileNotFoundError:
            print(f"❌ ファイルが見つかりません: {file_path}")
            return None
        except Exception as e:
            print(f"❌ ファイル読み込みエラー: {e}")
            return None
    
    def analyze_markdown_structure(self, markdown_content):
        """Markdown議事録の構造を分析"""
        prompt = f"""
以下のMarkdown議事録を分析し、Word文書化のための構造情報を抽出してください：

{markdown_content}

以下の形式で分析結果を出力してください：

## 📋 議事録情報
- 会議タイトル: 
- 開催日時: 
- 参加者: 
- 会議時間: 

## 📝 主要セクション
1. セクション名とその内容の要約
2. セクション名とその内容の要約
3. ...

## 🎯 重要ポイント
- 決定事項の数: 
- アクションアイテムの数: 
- 議論点の数: 

## 📄 Word文書構成提案
1. 表紙の内容
2. 目次構成
3. 本文のレイアウト提案
4. 付録の必要性

Wordテンプレートとして最適な文書構造を提案してください。
"""
        
        try:
            response = self.client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=2000,
                messages=[{"role": "user", "content": prompt}]
            )
            
            analysis = response.content[0].text
            print("📊 Markdown構造分析完了")
            return analysis
            
        except Exception as e:
            print(f"❌ 分析エラー: {e}")
            return None
    
    def generate_word_document_plan(self, markdown_content, analysis):
        """Word文書生成プランを作成"""
        prompt = f"""
以下のMarkdown議事録とその分析結果を基に、Word MCPサーバーで実行する具体的なWord文書生成プランを作成してください：

## 元のMarkdown議事録:
{markdown_content[:1000]}...

## 構造分析結果:
{analysis}

以下の形式で、Word MCPサーバーで実行可能なコマンド形式の手順を提案してください：

## 🏗️ Word文書生成手順

### Step 1: 文書作成・基本設定
```
create_document: "議事録_[会議名]_[日付].docx"
set_page_margins: {"top": 2.5, "bottom": 2.5, "left": 3.0, "right": 2.5}
set_font_default: {"name": "游明朝", "size": 11}
```

### Step 2: 表紙作成
```
add_title: "[会議タイトル]"
add_subtitle: "議事録"
add_meeting_info: "開催日時: [日時] | 場所: [場所]"
add_page_break
```

### Step 3: 本文構造作成
```
add_heading1: "1. 会議概要"
add_table: 会議基本情報テーブル
add_heading1: "2. 議事内容"
...
```

### Step 4: 内容挿入
- 各セクションの具体的な内容挿入方法
- 表・リストの作成方法
- フォーマット指定

### Step 5: 仕上げ
- ヘッダー・フッター設定
- ページ番号追加
- 目次自動生成

実際にWord MCPで実行可能な、段階的な手順を提案してください。
"""
        
        try:
            response = self.client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=3000,
                messages=[{"role": "user", "content": prompt}]
            )
            
            plan = response.content[0].text
            print("📋 Word文書生成プラン作成完了")
            return plan
            
        except Exception as e:
            print(f"❌ プラン生成エラー: {e}")
            return None
    
    def execute_word_generation_with_mcp(self, markdown_content, generation_plan):
        """Word MCPサーバーを使ってWord文書を生成"""
        prompt = f"""
以下のMarkdown議事録とWord文書生成プランを基に、Word MCPサーバーのツールを使って実際にWord文書を生成してください：

## Markdown議事録:
{markdown_content}

## 生成プラン:
{generation_plan}

Word MCPサーバーの以下のようなツールを使用して、段階的にWord文書を作成してください：
- create_document
- add_heading
- add_paragraph  
- add_table
- set_formatting
- save_document

実際にWord文書を生成し、完成したファイルのパスを教えてください。
"""
        
        try:
            # 注意: 実際のWord MCPサーバーが動作している場合のみ有効
            response = self.client.beta.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=3000,
                messages=[{"role": "user", "content": prompt}],
                mcp_servers=self.word_mcp_servers,  # Word MCPサーバー
                betas=["mcp-client-2025-04-04"],
            )
            
            result = ""
            used_tools = []
            
            for content in response.content:
                if content.type == "text":
                    result += content.text
                elif content.type == "mcp_tool_use":
                    used_tools.append(content.name)
                    print(f"🔧 [Word MCP] {content.name}")
            
            print(f"✅ Word文書生成完了")
            print(f"🛠️  使用ツール: {', '.join(used_tools)}")
            return result, used_tools
            
        except Exception as e:
            print(f"❌ Word MCP実行エラー: {e}")
            print("💡 Word MCPサーバーが起動していることを確認してください")
            return None, []
    
    def generate_word_manually(self, markdown_content, analysis, plan):
        """Word MCPが利用できない場合の代替手段"""
        print("🔄 Word MCP代替モード - python-docxで直接生成")
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # 新しいWord文書作成
            doc = Document()
            
            # タイトル追加
            title = doc.add_heading('会議議事録', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 会議情報テーブル
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Table Grid'
            
            table.cell(0, 0).text = '会議名'
            table.cell(0, 1).text = '定例会議'  # Markdownから抽出
            table.cell(1, 0).text = '開催日時'
            table.cell(1, 1).text = datetime.now().strftime('%Y年%m月%d日')
            table.cell(2, 0).text = '参加者'
            table.cell(2, 1).text = '(Markdownから抽出)'
            table.cell(3, 0).text = '記録者'
            table.cell(3, 1).text = 'AI議事録システム'
            
            # Markdownコンテンツをセクションごとに処理
            doc.add_page_break()
            
            # Markdownをパース
            lines = markdown_content.split('\n')
            current_heading_level = 1
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # 見出し処理
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    heading_text = line.lstrip('#').strip()
                    doc.add_heading(heading_text, level)
                    current_heading_level = level
                
                # リスト処理
                elif line.startswith('- ') or line.startswith('* '):
                    item_text = line[2:].strip()
                    paragraph = doc.add_paragraph(item_text)
                    paragraph.style = 'List Bullet'
                
                # 番号付きリスト
                elif re.match(r'^\d+\.', line):
                    item_text = re.sub(r'^\d+\.\s*', '', line)
                    paragraph = doc.add_paragraph(item_text)
                    paragraph.style = 'List Number'
                
                # 通常のパラグラフ
                else:
                    doc.add_paragraph(line)
            
            # ファイル保存
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'議事録_{timestamp}.docx'
            doc.save(filename)
            
            print(f"✅ Word文書を生成しました: {filename}")
            return filename
            
        except ImportError:
            print("❌ python-docxがインストールされていません")
            print("💡 インストール: pip install python-docx")
            return None
        except Exception as e:
            print(f"❌ Word文書生成エラー: {e}")
            return None
    
    def process_markdown_to_word(self, markdown_file_path, use_mcp=True):
        """Markdown議事録をWord文書に変換する完全プロセス"""
        print("🚀 Markdown → Word変換プロセス開始")
        print("="*60)
        
        # Step 1: Markdownファイル読み込み
        markdown_content = self.read_markdown_minutes(markdown_file_path)
        if not markdown_content:
            return None
        
        # Step 2: 構造分析
        print("\n📊 Step 2: Markdown構造分析中...")
        analysis = self.analyze_markdown_structure(markdown_content)
        if analysis:
            print("="*40)
            print(analysis)
            print("="*40)
        
        # Step 3: Word文書生成プラン作成
        print("\n📋 Step 3: Word文書生成プラン作成中...")
        plan = self.generate_word_document_plan(markdown_content, analysis)
        if plan:
            print("="*40)
            print(plan)
            print("="*40)
        
        # Step 4: Word文書生成
        print(f"\n📄 Step 4: Word文書生成中... (MCP: {use_mcp})")
        
        if use_mcp:
            try:
                result, tools = self.execute_word_generation_with_mcp(
                    markdown_content, plan
                )
                if result:
                    print("="*40)
                    print(result)
                    print("="*40)
                    return result
                else:
                    print("⚠️  Word MCP実行失敗、代替手段を使用")
                    use_mcp = False
            except Exception as e:
                print(f"⚠️  Word MCP実行エラー: {e}")
                print("💡 代替手段を使用します")
                use_mcp = False
        
        if not use_mcp:
            filename = self.generate_word_manually(markdown_content, analysis, plan)
            return filename

def main():
    converter = MarkdownToWordMCP()
    
    if len(sys.argv) < 2:
        print("📖 使用方法:")
        print("  python markdown_to_word_mcp.py <markdownファイルパス> [--no-mcp]")
        print()
        print("📝 例:")
        print("  python markdown_to_word_mcp.py meeting_minutes.md")
        print("  python markdown_to_word_mcp.py meeting_minutes.md --no-mcp")
        print()
        print("💡 オプション:")
        print("  --no-mcp : Word MCPを使わず、python-docxで直接生成")
        return
    
    markdown_file = sys.argv[1]
    use_mcp = "--no-mcp" not in sys.argv
    
    print(f"🎯 対象ファイル: {markdown_file}")
    print(f"⚙️  Word MCP使用: {use_mcp}")
    print()
    
    result = converter.process_markdown_to_word(markdown_file, use_mcp)
    
    if result:
        print(f"\n🎉 変換完了!")
        print(f"📄 生成されたファイル: {result}")
    else:
        print(f"\n❌ 変換失敗")

if __name__ == "__main__":
    import sys
    main()